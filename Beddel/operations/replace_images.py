# operations/replace_images.py

import hashlib
from docx.shared import Inches

def image_hash(image_bytes):
    return hashlib.sha256(image_bytes).hexdigest()

def process(doc, old_img_path, new_img_path):
    old_hash = image_hash(old_img_path.read_bytes())
    new_img_path_str = str(new_img_path)

    for rel in list(doc.part._rels.values()):
        if "image" in rel.reltype:
            img_data = rel.target_part.blob
            if image_hash(img_data) == old_hash:
                for p in doc.paragraphs:
                    if rel.rId in p._element.xml:
                        p.clear()
                        run = p.add_run()
                        run.add_picture(new_img_path_str, width=Inches(2.0))  # adjust size as needed
                        break
