# operations/merge_files.py

from docx import Document

def process(documents, out_doc):
    """
    Merge multiple documents into one with a TOC placeholder.

    Args:
        documents: List of pathlib.Path objects for input files
        out_doc: The target Document object to populate
    """

    out_doc.add_paragraph("Table of Contents", style="Title")
    out_doc.add_paragraph(
    "[To generate table of contents: \n"
    "- In Word: References > Table of Contents > Insert Automatic Table\n",
    style="Normal"
)



    for path in documents:
        # Heading for file name (used for TOC generation later)
        out_doc.add_paragraph(path.stem, style="Heading 1")

        doc = Document(path)
        for element in doc.element.body:
            out_doc.element.body.append(element)

        out_doc.add_paragraph()  # Space between documents
