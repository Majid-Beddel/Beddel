import re
from docx.shared import Pt, RGBColor

def process(doc, pattern, font_name=None, font_size=None, bold=None, italic=None, underline=None, color=None):
    print(f"📝 DEBUG: Running format_regex_matches")
    print(f"    Pattern: {pattern}")
    print(f"    Font name: {font_name}")
    print(f"    Font size: {font_size}")
    print(f"    Bold: {bold}")
    print(f"    Italic: {italic}")
    print(f"    Underline: {underline}")
    print(f"    Color: {color}")

    regex = re.compile(pattern)

    def format_run(run):
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = Pt(font_size)
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if underline is not None:
            run.font.underline = underline
        if color:
            try:
                rgb = RGBColor(*color)
                run.font.color.rgb = rgb
                print(f"    ✔️ Applied color {color} to run: '{run.text}'")
            except Exception as e:
                print(f"    ⚠️ ERROR applying color to run '{run.text}': {e}")

    # Apply formatting in paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if regex.search(run.text):
                print(f"✔️ Match found in paragraph run: '{run.text}'")
                format_run(run)

    # Apply formatting in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if regex.search(run.text):
                            print(f"✔️ Match found in table cell run: '{run.text}'")
                            format_run(run)
