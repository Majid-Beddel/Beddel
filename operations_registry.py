# from docx import Document
# import pathlib
# import dateparser
# from datetime import datetime

# from operations import (
#     bulk_replace,
#     format_headers,
#     format_body,
#     update_footers,
#     merge_files,
#     format_specific_words,
#     format_regex_matches,
#     regex_replace,
#     convert_to_pdf,
#     date_range_validator,
#     format_document,
#     table_cleanup,
#     extract_sections
# )

# # 🔧 Registry of available operations
# AVAILABLE_OPERATIONS = {
#     "bulk_replace": {
#         "description": "Bulk replace text in document",
#         "func": bulk_replace.process
#     },
#     "format_headers": {
#         "description": "Format specified headings (e.g. Heading 1, 2)",
#         "func": format_headers.process
#     },
#     "format_body": {
#         "description": "Format normal body text (non-headings)",
#         "func": format_body.process
#     },
#     "merge_files": {
#         "description": "Merge all files in folder into one document with TOC",
#         "func": merge_files.process
#     },
#     "format_specific_words": {
#         "description": "Format all occurrences of a specific word or phrase",
#         "func": format_specific_words.process
#     },
#     "update_footers": {
#         "description": "Change footer text",
#         "func": update_footers.process
#     },
#     "format_regex_matches": {
#         "description": "Format text matching a regex pattern",
#         "func": format_regex_matches.process
#     },
#     "regex_replace": {
#         "description": "Replace text matching regex pattern with repl",
#         "func": regex_replace.process
#     },
#     "convert_to_pdf": {
#         "description": "Convert DOCX files to PDF",
#         "func": convert_to_pdf.process
#     },
#     "date_range_validator": {
#         "description": "Check if dates are within a given range",
#         "func": date_range_validator.process
#     },
#     "format_document": {
#         "description": "Adjust layout: align text, tables, resize images",
#         "func": format_document.process
#     },
#     "table_cleanup": {
#         "description": "Clean and normalize all tables in the document",
#         "func": table_cleanup.process
#     },
#     "extract_sections": {
#         "description": "Extract a specific section by heading (e.g. About, Experience)",
#         "func": extract_sections.process
#     }
# }

# # 🔧 Shared utility used by process_plan.py:
# def run_operation(operation_name, src_dir, out_dir, params):
#     src_path = pathlib.Path(src_dir)
#     out_path = pathlib.Path(out_dir)
#     out_path.mkdir(parents=True, exist_ok=True)

#     docx_files = sorted(src_path.glob("*.docx"))
#     if not docx_files:
#         return {"status": "error", "message": f"No .docx files found in {src_dir}"}

#     if operation_name == "merge_files":
#         out_doc = Document()
#         merge_files.process(docx_files, out_doc)
#         out_file = out_path / "merged_files.docx"
#         out_doc.save(out_file)
#         return {"status": "success", "output": str(out_file)}

#     if operation_name == "date_range_validator":
#         import re
#         date_pattern = r"(\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\w+ \d{4}\b|\b\d{4}-\d{2}-\d{2}\b)"
#         start = datetime.strptime(params.get("start_date"), "%Y-%m-%d")
#         end = datetime.strptime(params.get("end_date"), "%Y-%m-%d")

#         report_doc = Document()
#         report_doc.add_heading("Date Range Validation Report", level=1)

#         for file_path in docx_files:
#             doc = Document(file_path)
#             report_doc.add_paragraph(f"📄 File: {file_path.name}")
#             out_of_range = []

#             for para in doc.paragraphs:
#                 matches = re.findall(date_pattern, para.text)
#                 for match in matches:
#                     parsed = dateparser.parse(match)
#                     if parsed:
#                         if parsed.date() < start.date() or parsed.date() > end.date():
#                             out_of_range.append((match, parsed.date()))

#             if out_of_range:
#                 for match, parsed_date in out_of_range:
#                     report_doc.add_paragraph(
#                         f" - ❌ '{match}' → {parsed_date} is OUTSIDE the range {start.date()} to {end.date()}"
#                     )
#             else:
#                 report_doc.add_paragraph(" - ✅ All dates within range.")

#             report_doc.add_paragraph("")  # spacing

#         report_path = out_path / "date_validation_report.docx"
#         report_doc.save(report_path)

#         return {
#             "status": "success",
#             "message": f"Validation complete. Report saved to {report_path.name}",
#             "report_file": str(report_path)
#         }

#     if operation_name not in AVAILABLE_OPERATIONS:
#         return {"status": "error", "message": f"Unsupported operation: {operation_name}"}

#     op_func = AVAILABLE_OPERATIONS[operation_name]["func"]

#     for file_path in docx_files:
#         output_path = out_path / file_path.name

#         # Remove conflicting args from params if present
#         clean_params = params.copy()
#         clean_params.pop("input_path", None)
#         clean_params.pop("output_path", None)

#         if operation_name == "convert_to_pdf":
#             output_path = out_path / file_path.with_suffix(".pdf").name
#             op_func(
#                 input_path=file_path,
#                 output_path=output_path,
#                 use_word=clean_params.get("use_word", False)
#             )
#         elif operation_name == "extract_sections":
#             op_func(
#                 input_path=file_path,
#                 out_dir=out_path,
#                 **clean_params
#             )
#         else:
#             op_func(
#                 input_path=file_path,
#                 output_path=output_path,
#                 **clean_params
#             )

#     return {
#         "status": "success",
#         "files_processed": len(docx_files),
#         "output_dir": str(out_path)
#     }

from docx import Document
from datetime import datetime
import dateparser
import re

from operations import (
    bulk_replace,
    format_headers,
    format_body,
    update_footers,
    merge_files,
    redact,
    format_specific_words,
    format_regex_matches,
    regex_replace,
    convert_to_pdf,
    date_range_validator,
    format_document,
    table_cleanup,
    extract_sections
    
)

# 🔧 Registry of available operations
AVAILABLE_OPERATIONS = {
<<<<<<< HEAD
    "bulk_replace": {"description": "Bulk replace text in document", "func": bulk_replace.process},
    "format_headers": {"description": "Format specified headings", "func": format_headers.process},
    "format_body": {"description": "Format body text", "func": format_body.process},
    "merge_files": {"description": "Merge all files into one", "func": merge_files.process},
    "format_specific_words": {"description": "Format specific words", "func": format_specific_words.process},
    "update_footers": {"description": "Change footer text", "func": update_footers.process},
    "format_regex_matches": {"description": "Format regex matches", "func": format_regex_matches.process},
    "regex_replace": {"description": "Regex text replace", "func": regex_replace.process},
    "convert_to_pdf": {"description": "Convert to PDF", "func": convert_to_pdf.process},
    "date_range_validator": {"description": "Validate date ranges", "func": date_range_validator.process},
    "format_document": {"description": "Document layout cleanup", "func": format_document.process},
    "table_cleanup": {"description": "Clean tables", "func": table_cleanup.process},
    "extract_sections": {"description": "Extract section(s) by heading", "func": extract_sections.process},
}

=======
    "bulk_replace": {
        "description": "Bulk replace text in document",
        "func": bulk_replace.process
    },
    "redact": {
    "description": "Redact sensitive text using regex",
    "func": redact.process
    },
    "format_headers": {
        "description": "Format specified headings (e.g. Heading 1, 2)",
        "func": format_headers.process
    },
    "format_body": {
        "description": "Format normal body text (non-headings)",
        "func": format_body.process
    },
    "merge_files": {
        "description": "Merge all files in folder into one document with TOC",
        "func": merge_files.process
    },
    "format_specific_words": {
        "description": "Format all occurrences of a specific word or phrase",
        "func": format_specific_words.process
    },
    "update_footers": {
        "description": "Change footer text",
        "func": update_footers.process
    },
    "format_regex_matches": {
        "description": "Format text matching a regex pattern",
        "func": format_regex_matches.process
    },
    "regex_replace": {
        "description": "Replace text matching regex pattern with repl",
        "func": regex_replace.process
    },
    "convert_to_pdf": {
        "description": "Convert DOCX files to PDF",
        "func": convert_to_pdf.process
    },
    "date_range_validator": {
        "description": "Check if dates are within a given range",
        "func": date_range_validator.process
    },
    "format_document": {
        "description": "Adjust layout: align text, tables, resize images",
        "func": format_document.process
    },
    "table_cleanup": {
        "description": "Clean and normalize all tables in the document",
        "func": table_cleanup.process
    },
    "extract_sections": {
        "description": "Extract a specific section by heading (e.g. About, Experience)",
        "func": extract_sections.process
    }
}

# 🔧 Shared utility used by process_plan.py:
def run_operation(operation_name, src_dir, out_dir, params):
    src_path = pathlib.Path(src_dir)
    out_path = pathlib.Path(out_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    docx_files = sorted(src_path.glob("*.docx"))
    if not docx_files:
        return {"status": "error", "message": f"No .docx files found in {src_dir}"}

    if operation_name == "merge_files":
        out_doc = Document()
        merge_files.process(docx_files, out_doc)
        out_file = out_path / "merged_files.docx"
        out_doc.save(out_file)
        return {"status": "success", "output": str(out_file)}

    if operation_name == "date_range_validator":
        import re
        date_pattern = r"(\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\w+ \d{4}\b|\b\d{4}-\d{2}-\d{2}\b)"
        start = datetime.strptime(params.get("start_date"), "%Y-%m-%d")
        end = datetime.strptime(params.get("end_date"), "%Y-%m-%d")

        report_doc = Document()
        report_doc.add_heading("Date Range Validation Report", level=1)

        for file_path in docx_files:
            doc = Document(file_path)
            report_doc.add_paragraph(f"📄 File: {file_path.name}")
            out_of_range = []

            for para in doc.paragraphs:
                matches = re.findall(date_pattern, para.text)
                for match in matches:
                    parsed = dateparser.parse(match)
                    if parsed:
                        if parsed.date() < start.date() or parsed.date() > end.date():
                            out_of_range.append((match, parsed.date()))

            if out_of_range:
                for match, parsed_date in out_of_range:
                    report_doc.add_paragraph(
                        f" - ❌ '{match}' → {parsed_date} is OUTSIDE the range {start.date()} to {end.date()}"
                    )
            else:
                report_doc.add_paragraph(" - ✅ All dates within range.")
            report_doc.add_paragraph("")  # spacing

        report_path = out_path / "date_validation_report.docx"
        report_doc.save(report_path)

        return {
            "status": "success",
            "message": f"Validation complete. Report saved to {report_path.name}",
            "report_file": str(report_path)
        }
>>>>>>> 2741e0a (Redline feature complete: regex_replace, table_cleanup, etc.)

def run_operation(operation_name, doc, params):
    if operation_name not in AVAILABLE_OPERATIONS:
        return {"status": "error", "message": f"Unsupported operation: {operation_name}"}

    op_func = AVAILABLE_OPERATIONS[operation_name]["func"]

<<<<<<< HEAD
    try:
        result = op_func(doc, **params)
        return {"status": "success", "result": result}
    except Exception as e:
        return {"status": "error", "message": str(e)}
=======
    for file_path in docx_files:
        output_path = out_path / file_path.name

        # Clean incoming params
        clean_params = params.copy()
        clean_params.pop("input_path", None)
        clean_params.pop("output_path", None)

        print(f"🚀 Running operation '{operation_name}' on {file_path.name}")
        print(f"📦 Params: {clean_params}")

        try:
            if operation_name == "convert_to_pdf":
                output_path = out_path / file_path.with_suffix(".pdf").name
                op_func(
                    input_path=file_path,
                    output_path=output_path,
                    use_word=clean_params.get("use_word", False)
                )
            elif operation_name == "extract_sections":
                op_func(
                    input_path=file_path,
                    out_dir=out_path,
                    **clean_params
                )
            else:
                op_func(
                    input_path=file_path,
                    output_path=output_path,
                    **clean_params
                )
        except Exception as e:
            print(f"❌ Error during '{operation_name}': {type(e).__name__} - {e}")
            import traceback
            traceback.print_exc()
            return {
                "status": "error",
                "message": f"Operation '{operation_name}' failed on {file_path.name}"
            }

    # ✅ Check output files created
    output_files = list(out_path.glob("*.docx")) + list(out_path.glob("*.pdf"))

    if not output_files:
        return {
            "status": "error",
            "message": f"No output files created in {out_path}"
        }

    print("📄 Output files created:")
    for f in output_files:
        print(f" - {f.name} ({f.stat().st_size} bytes)")

    return {
        "status": "success",
        "files_processed": len(docx_files),
        "output_files": [str(f) for f in output_files],
        "output_dir": str(out_path)
    }
>>>>>>> 2741e0a (Redline feature complete: regex_replace, table_cleanup, etc.)
