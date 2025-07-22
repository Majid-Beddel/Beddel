from docx import Document

def process(doc, input_path=None, out_dir=None, target_heading=None):
    section_doc = Document()
    capture = False

    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            if target_heading.lower() in para.text.strip().lower():
                capture = True
                section_doc.add_paragraph(para.text, style=para.style)
            else:
                capture = False
        elif capture:
            section_doc.add_paragraph(para.text)

    if section_doc.paragraphs:
        if input_path and out_dir:
            output_filename = f"{input_path.stem}_{target_heading.lower().replace(' ', '_')}.docx"
            output_path = out_dir / output_filename
            section_doc.save(output_path)
            print(f"✅ Extracted section saved to: {output_path}")
        else:
            print("⚠️ Missing input_path or out_dir — cannot save extracted section.")
    else:
        name = input_path.name if input_path else "unknown"
        print(f"⚠️ Section '{target_heading}' not found in {name}")