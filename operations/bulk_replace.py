from docx import Document
from docx.shared import RGBColor
from pathlib import Path
import re

def replace_case_insensitive_split(text, find, repl):
    pattern = re.compile(re.escape(find), re.IGNORECASE)
    result = []
    last_end = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_end:
            result.append((text[last_end:start], False))  # Unchanged
        original = match.group()
        if original.isupper():
            repl_text = repl.upper()
        elif original[0].isupper():
            repl_text = repl.capitalize()
        else:
            repl_text = repl
        result.append((repl_text, True))  # Replacement
        last_end = end
    if last_end < len(text):
        result.append((text[last_end:], False))
    return result

def process(input_path, output_path, find=None, repl=None, redline=False):
    if not find or repl is None:
        return

    doc = Document(input_path)
    replacement_count = 0

    for paragraph in doc.paragraphs:
        original_runs = list(paragraph.runs)  # Save original runs
        full_text = ''.join(run.text for run in original_runs)

        if re.search(re.escape(find), full_text, flags=re.IGNORECASE):
            # Clear paragraph runs
            for run in original_runs:
                run.clear()

            new_parts = replace_case_insensitive_split(full_text, find, repl)
            for text, is_replacement in new_parts:
                if not text:
                    continue
                new_run = paragraph.add_run(text)
                if is_replacement and redline:
                    new_run.font.color.rgb = RGBColor(255, 0, 0)
                    new_run.bold = True
                replacement_count += is_replacement

    doc.save(output_path)

    # Create redline version
    if redline:
        output_path_obj = Path(output_path)
        redline_path = output_path_obj.parent / f"redline_{output_path_obj.name}"

        redline_doc = Document(output_path)
        summary_paragraph = redline_doc.add_paragraph()
        summary_paragraph.add_run(
            f"🔄 Redline Summary: Replaced '{find}' → '{repl}' ({replacement_count} change{'s' if replacement_count != 1 else ''})"
        ).bold = True

        # Insert summary at the top
        body = redline_doc._body._element
        body.insert(0, summary_paragraph._element)

        redline_doc.save(redline_path)
