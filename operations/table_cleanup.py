# operations/table_cleanup.py

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import RGBColor
from pathlib import Path

def process(input_path, output_path, redline=False):
    doc = Document(input_path)
    cleaned_cells = 0

    def cleanup_tables(doc, highlight=False):
        nonlocal cleaned_cells

        for table in doc.tables:
            num_cols = len(table.columns)
            even_width = 10000 // num_cols if num_cols else 2500

            for row in table.rows:
                tr = row._tr
                trPr = tr.get_or_add_trPr()

                # Set fixed row height
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), '500')
                trHeight.set(qn('w:hRule'), 'exact')
                trPr.append(trHeight)

                for cell in row.cells:
                    # Normalize alignment + strip whitespace
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cleaned_cells += 1
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para.text = para.text.strip()

                        if highlight:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 0, 0)

                    # Normalize width
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w = OxmlElement('w:tcW')
                    tc_w.set(qn('w:type'), 'dxa')
                    tc_w.set(qn('w:w'), str(even_width))
                    tc_pr.append(tc_w)

                    # Remove broken merges
                    v_merge = tc_pr.find(qn('w:vMerge'))
                    if v_merge is not None:
                        tc_pr.remove(v_merge)

    # Process the main document
    cleanup_tables(doc)
    doc.save(output_path)

    if redline:
        red_doc = Document(input_path)
        cleaned_cells = 0  # reset counter for redline version
        cleanup_tables(red_doc, highlight=True)

        # Add redline summary
        summary = f"🔄 Redline Summary:\nCleaned and formatted {cleaned_cells} cell(s) with even widths, alignment, and removed broken merges"
        summary_para = red_doc.add_paragraph(summary)
        first_para = red_doc.paragraphs[0]
        first_para._element.addprevious(summary_para._element)

        redline_path = Path(output_path).parent / f"redline_{Path(output_path).name}"
        red_doc.save(redline_path)
