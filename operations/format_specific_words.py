# operations/format_specific_words.py

from docx.shared import Pt, RGBColor
<<<<<<< HEAD
=======
from pathlib import Path

def process(input_path, output_path, keyword=None, font_name=None, font_size=None, bold=None, italic=None, underline=None, color=None, redline=False):
    doc = Document(input_path)
>>>>>>> 2741e0a (Redline feature complete: regex_replace, table_cleanup, etc.)

def process(doc, keyword=None, font_name=None, font_size=None, bold=None, italic=None, underline=None, color=None):
    if not keyword:
        return

    keyword_lower = keyword.lower()
    keyword_len = len(keyword)
    count = 0  # Count number of keyword matches

    for paragraph in doc.paragraphs:
        new_runs = []

        for run in paragraph.runs:
            text = run.text
            text_lower = text.lower()

            idx = 0
            while idx < len(text):
                match_idx = text_lower.find(keyword_lower, idx)
                if match_idx == -1:
                    # No more matches, add remainder as normal run
                    remainder = text[idx:]
                    if remainder:
                        new_run = paragraph.add_run(remainder)
                        new_runs.append(new_run)
                    break
                else:
                    # Add text before match
                    if match_idx > idx:
                        before = text[idx:match_idx]
                        new_run = paragraph.add_run(before)
                        new_runs.append(new_run)

                    # Add matched keyword with formatting
                    matched = text[match_idx:match_idx + keyword_len]
                    keyword_run = paragraph.add_run(matched)

                    if font_name:
                        keyword_run.font.name = font_name
                    if font_size:
                        keyword_run.font.size = Pt(font_size)
                    if bold is not None:
                        keyword_run.font.bold = bold
                    if italic is not None:
                        keyword_run.font.italic = italic
                    if underline is not None:
                        keyword_run.font.underline = underline
                    if color:
                        keyword_run.font.color.rgb = RGBColor(*color)

                    new_runs.append(keyword_run)
                    count += 1
                    idx = match_idx + keyword_len

        # After building all new runs, remove old runs
        for old_run in paragraph.runs:
            p = old_run._element
            p.getparent().remove(p)

        # Add new runs back in order
        for r in new_runs:
            paragraph._element.append(r._element)
<<<<<<< HEAD
=======

    doc.save(output_path)

    # ✅ Redline logic
    if redline:
        output_path_obj = Path(output_path)
        redline_path = output_path_obj.parent / f"redline_{output_path_obj.name}"

        redline_doc = Document(output_path)

        # Redline formatting pass (same text, but force red + bold)
        for para in redline_doc.paragraphs:
            for run in para.runs:
                if keyword_lower in run.text.lower():
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)

        # Add summary paragraph at the very top via XML
        summary_text = f"🔄 Redline Summary:\nFormatted keyword '{keyword}' {count} time(s)"
        if bold is not None:
            summary_text += f" with bold={bold}"
        if italic is not None:
            summary_text += f", italic={italic}"
        if underline is not None:
            summary_text += f", underline={underline}"
        if color:
            summary_text += f", color=RGB{tuple(color)}"

        para = redline_doc.add_paragraph(summary_text)
        redline_doc._body._element.insert(0, para._element)

        redline_doc.save(redline_path)
>>>>>>> 2741e0a (Redline feature complete: regex_replace, table_cleanup, etc.)
