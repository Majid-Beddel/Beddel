import re
from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path

def process(input_path, output_path, pattern=None, font_name=None, font_size=None,
            bold=None, italic=None, underline=None, color=None, redline=False):
    if not pattern:
        return

    regex = re.compile(pattern)
    match_count = 0

    # Map known patterns to friendly descriptions
    pattern_descriptions = {
        r"(\+44\s?7\d{3}|\(?07\d{3}\)?)\s?\d{3}\s?\d{3}": "UK phone number",
        # Add more here as needed
    }
    pattern_label = pattern_descriptions.get(pattern, f"pattern: '{pattern}'")

    def format_match_text(paragraph, text, match_spans, formatting_fn):
        runs = []
        idx = 0
        for start, end in match_spans:
            if start > idx:
                runs.append((text[idx:start], False))
            runs.append((text[start:end], True))
            idx = end
        if idx < len(text):
            runs.append((text[idx:], False))

        paragraph.clear()
        for part, is_match in runs:
            run = paragraph.add_run(part)
            if is_match:
                formatting_fn(run)

    def apply_formatting(run):
        nonlocal match_count
        match_count += 1
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = Pt(font_size)
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if underline is not None:
            run.font.underline = underline
        if color:
            run.font.color.rgb = RGBColor(*color)

    def apply_redline(run):
        nonlocal match_count
        match_count += 1
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)

    # Main document formatting
    doc = Document(input_path)
    for para in doc.paragraphs:
        text = para.text
        matches = list(regex.finditer(text))
        if matches:
            spans = [(m.start(), m.end()) for m in matches]
            format_match_text(para, text, spans, apply_formatting)

    doc.save(output_path)

    # Redline version
    if redline:
        redline_path = Path(output_path).parent / f"redline_{Path(output_path).name}"
        red_doc = Document(input_path)
        redline_count = 0  # Independent count

        for para in red_doc.paragraphs:
            text = para.text
            matches = list(regex.finditer(text))
            if matches:
                spans = [(m.start(), m.end()) for m in matches]
                format_match_text(para, text, spans, apply_redline)
                redline_count += len(matches)

        # Friendly summary at the top
        summary_text = f"🔄 Redline Summary:\nFormatted {redline_count} match(es) for {pattern_label}"
        first_para = red_doc.paragraphs[0]
        first_para.insert_paragraph_before(summary_text)

        red_doc.save(redline_path)
