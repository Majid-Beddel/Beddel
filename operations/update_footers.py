from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx import Document
from pathlib import Path

def process(input_path, output_path, footer_text=None, image_path=None, image_width_in_inches=2.0, redline=False):
    doc = Document(input_path)

    # Enable "Different First Page" so footer only appears on page 1
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.first_page_footer

    # Clear any existing content in footer
    for paragraph in footer.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Add footer text if provided
    if footer_text:
        para = footer.add_paragraph(footer_text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Add image if provided
    if image_path:
        try:
            img_para = footer.add_paragraph()
            run = img_para.add_run()
            run.add_picture(image_path, width=Inches(image_width_in_inches))
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"⚠️ Could not add image: {e}")

    doc.save(output_path)

    # Handle redline
    if redline and footer_text:
        output_path_obj = Path(output_path)
        redline_path = output_path_obj.parent / f"redline_{output_path_obj.name}"
        red_doc = Document(input_path)

        section = red_doc.sections[0]
        section.different_first_page_header_footer = True
        footer = section.first_page_footer

        for paragraph in footer.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)

        red_para = footer.add_paragraph(footer_text)
        red_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = red_para.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        if image_path:
            try:
                img_para = footer.add_paragraph()
                run = img_para.add_run()
                run.add_picture(image_path, width=Inches(image_width_in_inches))
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"⚠️ Could not add image to redline footer: {e}")

        # Add summary at the top of the document
        summary_text = f"🔄 Redline Summary:\nUpdated footer with text: \"{footer_text}\" (First page only)"
        summary_para = red_doc.add_paragraph(summary_text)
        red_doc.paragraphs[0]._element.addprevious(summary_para._element)

        red_doc.save(redline_path)
