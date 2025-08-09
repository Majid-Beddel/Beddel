# operations/redact.py

import re
from docx import Document
from docx.shared import RGBColor

def process(input_path, output_path, pattern, redline=False):
    regex = re.compile(pattern)
    doc = Document(input_path)

    redline_doc = Document()
    redline_summary = redline_doc.add_paragraph("🔍 Summary of Redactions:\n")

    num_replacements = 0

    for para in doc.paragraphs:
        if regex.search(para.text):
            para.text = regex.sub("[REDACTED]", para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if regex.search(cell.text):
                    cell.text = regex.sub("[REDACTED]", cell.text)

    doc.save(output_path)

    if redline:
        source = Document(input_path)
        for para in source.paragraphs:
            new_para = redline_doc.add_paragraph()
            last_idx = 0
            for match in regex.finditer(para.text):
                new_para.add_run(para.text[last_idx:match.start()])
                red_run = new_para.add_run("[REDACTED]")
                red_run.font.color.rgb = RGBColor(255, 0, 0)
                red_run.bold = True
                num_replacements += 1
                last_idx = match.end()
            new_para.add_run(para.text[last_idx:])

        for table in source.tables:
            new_table = redline_doc.add_table(rows=0, cols=len(table.columns))
            for row in table.rows:
                row_cells = new_table.add_row().cells
                for i, cell in enumerate(row.cells):
                    para = cell.text
                    red_text = ""
                    last_idx = 0
                    red_cell = ""
                    for match in regex.finditer(para):
                        red_text += para[last_idx:match.start()]
                        red_text += "[REDACTED]"
                        last_idx = match.end()
                        num_replacements += 1
                    red_text += para[last_idx:]
                    row_cells[i].text = red_text

        redline_summary.add_run(f"\nTotal redactions: {num_replacements}")
        redline_path = output_path.with_name(f"redline_{output_path.name}")
        redline_doc.save(redline_path)
