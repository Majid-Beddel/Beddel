from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def process(doc, alignment="left", **kwargs):
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT
    }
    align_value = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # Align paragraphs (not headings)
    for para in doc.paragraphs:
        if not para.style.name.startswith("Heading"):
            para.alignment = align_value

    # Align text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = align_value

    # Resize inline images
    for shape in doc.inline_shapes:
        if shape.width > Inches(6):
            shape.width = Inches(6)
        if shape.height > Inches(8):
            shape.height = Inches(8)