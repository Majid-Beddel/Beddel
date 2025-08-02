# import re
# from docx import Document
# from docx.shared import RGBColor
# from pathlib import Path

# def process(input_path, output_path, pattern=None, repl=None, redline=False):
#     if not pattern or repl is None:
#         return

#     regex = re.compile(pattern, flags=re.IGNORECASE)
#     replacement_count = 0

#     def replace_and_count(text):
#         nonlocal replacement_count
#         return regex.sub(lambda m: count(m, repl), text)

#     def count(match, repl):
#         nonlocal replacement_count
#         replacement_count += 1
#         return repl

#     # Process main document
#     doc = Document(input_path)

#     for para in doc.paragraphs:
#         para.text = replace_and_count(para.text)

#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 cell.text = replace_and_count(cell.text)

#     doc.save(output_path)

#     # ✅ Generate redline document
#     if redline:
#         redline_doc = Document(input_path)
#         redline_repl_count = 0

#         def format_redline(text):
#             matches = list(regex.finditer(text))
#             if not matches:
#                 return [(text, False)]

#             segments = []
#             last_idx = 0
#             for match in matches:
#                 if match.start() > last_idx:
#                     segments.append((text[last_idx:match.start()], False))
#                 segments.append((repl, True))
#                 last_idx = match.end()
#             if last_idx < len(text):
#                 segments.append((text[last_idx:], False))
#             return segments

#         for para in redline_doc.paragraphs:
#             parts = format_redline(para.text)
#             if any(is_red for _, is_red in parts):
#                 para.clear()
#                 for part, is_red in parts:
#                     run = para.add_run(part)
#                     if is_red:
#                         run.font.bold = True
#                         run.font.color.rgb = RGBColor(255, 0, 0)
#                         redline_repl_count += 1

#         for table in redline_doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     for para in cell.paragraphs:
#                         parts = format_redline(para.text)
#                         if any(is_red for _, is_red in parts):
#                             para.clear()
#                             for part, is_red in parts:
#                                 run = para.add_run(part)
#                                 if is_red:
#                                     run.font.bold = True
#                                     run.font.color.rgb = RGBColor(255, 0, 0)
#                                     redline_repl_count += 1

#         # ✅ Add summary paragraph at the very top
#         summary_text = f"🔄 Redline Summary:\nReplaced {redline_repl_count} occurrence(s) of pattern: '{pattern}' → '{repl}'"
#         summary_paragraph = redline_doc.paragraphs[0].insert_paragraph_before(summary_text)

#         redline_path = Path(output_path).parent / f"redline_{Path(output_path).name}"
#         redline_doc.save(redline_path)

import re
from docx import Document
from docx.shared import RGBColor
from pathlib import Path
import time

def process(input_path, output_path, pattern=None, repl=None, redline=False):
    if not pattern or repl is None:
        print("❌ Missing pattern or replacement.")
        return

    regex = re.compile(pattern, flags=re.IGNORECASE)
    replacement_count = 0

    def replace_and_count(text):
        nonlocal replacement_count
        return regex.sub(lambda m: count(m, repl), text)

    def count(match, repl):
        nonlocal replacement_count
        replacement_count += 1
        return repl

    # Process regular document
    doc = Document(input_path)

    for para in doc.paragraphs:
        para.text = replace_and_count(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_and_count(cell.text)

    doc.save(output_path)
    time.sleep(0.2)  # Give filesystem time to flush
    print(f"✅ Saved main output: {output_path} ({Path(output_path).stat().st_size} bytes)")

    # Now build redline version
    if redline:
        redline_doc = Document(input_path)
        redline_repl_count = 0

        def format_redline(text):
            matches = list(regex.finditer(text))
            if not matches:
                return [(text, False)]

            segments = []
            last_idx = 0
            for match in matches:
                if match.start() > last_idx:
                    segments.append((text[last_idx:match.start()], False))
                segments.append((repl, True))
                last_idx = match.end()
            if last_idx < len(text):
                segments.append((text[last_idx:], False))
            return segments

        def update_paragraph(para):
            nonlocal redline_repl_count
            parts = format_redline(para.text)
            para.clear()
            for part, is_red in parts:
                run = para.add_run(part)
                if is_red:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    redline_repl_count += 1

        for para in redline_doc.paragraphs:
            update_paragraph(para)

        for table in redline_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        update_paragraph(para)

        # Add summary paragraph
        summary_text = f"🔄 Redline Summary:\nReplaced {redline_repl_count} occurrence(s) of pattern: '{pattern}' → '{repl}'"
        summary_paragraph = redline_doc.paragraphs[0].insert_paragraph_before(summary_text)
        summary_paragraph.runs[0].bold = True

        redline_path = Path(output_path).parent / f"redline_{Path(output_path).name}"
        redline_doc.save(redline_path)
        time.sleep(0.2)
        print(f"✅ Saved redline file: {redline_path} ({redline_path.stat().st_size} bytes)")

