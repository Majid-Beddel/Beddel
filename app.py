# master app.py

import pathlib
import time
from tqdm import tqdm
from docx import Document

from operations import (
    bulk_replace,
    format_headers,
    format_body,
    update_footers,
    merge_files,
    format_specific_words
)


# Registry of available operations
AVAILABLE_OPERATIONS = {
    "bulk_replace": {
        "description": "Bulk replace text in document",
        "func": bulk_replace.process
    },
    "format_headers": {
        "description": "Format specified headings (e.g. Heading 1, 2)",
        "func": format_headers.process
    },
    "format_body": {
        "description": "Format normal body text (non-headings)",
        "func": format_body.process
    },
    "merge_files": {
        "description": "Merge all files in folder into one document with TOC",
        "func": merge_files.process
    },
    "format_specific_words": {
        "description": "Format all occurrences of a specific word or phrase",
        "func": format_specific_words.process
    },
    "update_footers": {
        "description": "Change footer text",
        "func": update_footers.process
    }
}

def prompt_directory(prompt_text):
    while True:
        path_str = input(prompt_text).strip()
        path = pathlib.Path(path_str)
        if path.exists() and path.is_dir():
            return path
        else:
            print(f"❌ '{path_str}' is not a valid directory. Please try again.")

def prompt_text(prompt_text):
    return input(prompt_text).strip()

def prompt_int(prompt_text):
    while True:
        try:
            value = int(input(prompt_text).strip())
            return value
        except ValueError:
            print("❌ Please enter a valid integer.")

def main():
    print("🔧 Document Processor CLI 🔧\n")

    # Dynamically list available operations
    print("Choose an operation:")
    op_keys = list(AVAILABLE_OPERATIONS.keys())
    for idx, key in enumerate(op_keys, 1):
        desc = AVAILABLE_OPERATIONS[key]["description"]
        print(f"{idx}. {key} - {desc}")

    while True:
        choice = input("> ").strip()
        if choice.isdigit() and 1 <= int(choice) <= len(op_keys):
            selected_key = op_keys[int(choice) - 1]
            break
        else:
            print(f"❌ Please enter a number between 1 and {len(op_keys)}.")

    src_dir = prompt_directory("\nWhat directory contains the files you want to process?\n> ")
    out_dir = prompt_directory("\nWhere should the updated files be saved?\n> ")
    out_dir.mkdir(parents=True, exist_ok=True)

    docx_files = sorted(src_dir.glob("*.docx"))
    if not docx_files:
        print(f"⚠️ No .docx files found in {src_dir}")
        return

    op_params = {}

    if selected_key == "bulk_replace":
        op_params["find"] = prompt_text("\nWhat is the word you want to replace?\n> ")
        op_params["repl"] = prompt_text("What is the replacement word?\n> ")

    elif selected_key == "format_headers":
        font_name = prompt_text("\nFont name (optional, press Enter to skip):\n> ") or None
        font_size = prompt_text("Font size (optional, press Enter to skip):\n> ")
        font_size = int(font_size) if font_size else None

        bold = prompt_text("Bold? (y/n, Enter to skip):\n> ")
        bold = True if bold.lower() == 'y' else False if bold.lower() == 'n' else None

        italic = prompt_text("Italic? (y/n, Enter to skip):\n> ")
        italic = True if italic.lower() == 'y' else False if italic.lower() == 'n' else None

        underline = prompt_text("Underline? (y/n, Enter to skip):\n> ")
        underline = True if underline.lower() == 'y' else False if underline.lower() == 'n' else None

        color_str = prompt_text("Color (e.g. 255,0,0 for red) or Enter to skip:\n> ")
        color = tuple(map(int, color_str.split(','))) if color_str else None

        levels = prompt_text("Heading levels (e.g. 1,2 or Enter for 1,2,3):\n> ")
        heading_levels = [int(l.strip()) for l in levels.split(',')] if levels else [1, 2, 3]

        op_params.update({
            "font_name": font_name,
            "font_size": font_size,
            "bold": bold,
            "italic": italic,
            "underline": underline,
            "color": color,
            "heading_levels": heading_levels
        })

    elif selected_key == "format_body":
        font_name = prompt_text("\nFont name (optional, press Enter to skip):\n> ") or None
        font_size = prompt_text("Font size (optional, press Enter to skip):\n> ")
        font_size = int(font_size) if font_size else None

        bold = prompt_text("Bold? (y/n, Enter to skip):\n> ")
        bold = True if bold.lower() == 'y' else False if bold.lower() == 'n' else None

        italic = prompt_text("Italic? (y/n, Enter to skip):\n> ")
        italic = True if italic.lower() == 'y' else False if italic.lower() == 'n' else None

        underline = prompt_text("Underline? (y/n, Enter to skip):\n> ")
        underline = True if underline.lower() == 'y' else False if underline.lower() == 'n' else None

        color_str = prompt_text("Color (e.g. 255,0,0 for red) or Enter to skip:\n> ")
        color = tuple(map(int, color_str.split(','))) if color_str else None

        op_params.update({
            "font_name": font_name,
            "font_size": font_size,
            "bold": bold,
            "italic": italic,
            "underline": underline,
            "color": color
        })

    elif selected_key == "insert_footer":
        footer_text = prompt_text("\nEnter footer text (optional, press Enter to skip):\n> ") or None
        image_path = prompt_text("Path to signature image (optional, press Enter to skip):\n> ") or None
        image_width = prompt_text("Image width in inches (optional, default 2.0):\n> ")
        image_width = float(image_width) if image_width else 2.0

        op_params.update({
            "footer_text": footer_text,
            "image_path": image_path if image_path else None,
            "image_width_in_inches": image_width
        })

    elif selected_key == "format_specific_words":
        keyword = prompt_text("\nWhich word or phrase should be formatted?\n> ")

        font_name = prompt_text("Font name (optional, Enter to skip):\n> ") or None
        font_size = prompt_text("Font size (optional, Enter to skip):\n> ")
        font_size = int(font_size) if font_size else None

        bold = prompt_text("Bold? (y/n, Enter to skip):\n> ")
        bold = True if bold.lower() == 'y' else False if bold.lower() == 'n' else None

        italic = prompt_text("Italic? (y/n, Enter to skip):\n> ")
        italic = True if italic.lower() == 'y' else False if italic.lower() == 'n' else None

        underline = prompt_text("Underline? (y/n, Enter to skip):\n> ")
        underline = True if underline.lower() == 'y' else False if underline.lower() == 'n' else None

        color_str = prompt_text("Color (e.g. 255,0,0 for red) or Enter to skip:\n> ")
        color = tuple(map(int, color_str.split(','))) if color_str else None

        op_params.update({
            "keyword": keyword,
            "font_name": font_name,
            "font_size": font_size,
            "bold": bold,
            "italic": italic,
            "underline": underline,
            "color": color
        })

    if selected_key == "merge_files":
        out_path = out_dir / "merged_files.docx"
        out_doc = Document()

        merge_files.process(docx_files, out_doc)
        out_doc.save(out_path)

        print(f"\n✅ Merged document saved as: {out_path}")
        return

    op_func = AVAILABLE_OPERATIONS[selected_key]["func"]

    print(f"\n✅ Found {len(docx_files)} files. Starting '{selected_key}'...\n")
    start_time = time.time()

    for file_path in tqdm(docx_files, desc=f"Processing files ({selected_key})"):
        doc = Document(file_path)
        op_func(doc, **op_params)
        out_path = out_dir / file_path.name
        doc.save(out_path)

    elapsed = time.time() - start_time
    print(f"\n✅ Completed {len(docx_files)} files in {elapsed:.2f} seconds.")
    print(f"👉 Output files saved to: {out_dir}")

def run_operation(operation_name, src_dir, out_dir, params):
    src_path = pathlib.Path(src_dir)
    out_path = pathlib.Path(out_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    docx_files = sorted(src_path.glob("*.docx"))
    if not docx_files:
        return {"status": "error", "message": f"No .docx files found in {src_dir}"}

    if operation_name == "merge_files":
        out_doc = Document()
        merge_files.process(docx_files, out_doc)
        out_file = out_path / "merged_files.docx"
        out_doc.save(out_file)
        return {"status": "success", "output": str(out_file)}

    if operation_name not in AVAILABLE_OPERATIONS:
        return {"status": "error", "message": f"Unsupported operation: {operation_name}"}

    op_func = AVAILABLE_OPERATIONS[operation_name]["func"]

    for file_path in docx_files:
        doc = Document(file_path)
        op_func(doc, **params)
        out_file = out_path / file_path.name
        doc.save(out_file)

    return {"status": "success", "files_processed": len(docx_files), "output_dir": str(out_path)}
    

if __name__ == "__main__":
    main()
